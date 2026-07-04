from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from packages.knowledge_store.models import KnowledgeItem


class TestParsePDF:
    """PDF 解析器单元测试"""

    def _make_test_pdf(
        self, path: Path, lines: list[str] | None = None
    ) -> None:
        """Helper: create a minimal PDF with extractable text.

        Uses insert_htmlbox to ensure CJK characters are properly embedded
        and extractable via get_text().
        """
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        if lines:
            html_content = "<br/>".join(f"<p>{line}</p>" for line in lines)
            rect = fitz.Rect(50, 50, 500, 50 + len(lines) * 50)
            page.insert_htmlbox(rect, html_content)
        doc.save(str(path))
        doc.close()

    def test_parse_pdf_returns_text(self, tmp_path: Path) -> None:
        """PDF 解析应返回纯文本"""
        pdf_path = tmp_path / "test.pdf"
        self._make_test_pdf(
            pdf_path, ["羊肚菌是一种珍贵的食用菌", "富含多种氨基酸和微量元素"]
        )

        from packages.knowledge_store.parsers import parse_pdf

        text = parse_pdf(pdf_path)
        assert "羊肚菌" in text
        assert "氨基酸" in text

    def test_parse_pdf_empty_document(self, tmp_path: Path) -> None:
        """空 PDF 应返回空字符串"""
        pdf_path = tmp_path / "empty.pdf"
        self._make_test_pdf(pdf_path)

        from packages.knowledge_store.parsers import parse_pdf

        text = parse_pdf(pdf_path)
        assert text.strip() == ""

    def test_parse_pdf_file_not_found(self) -> None:
        """不存在的 PDF 应抛出异常"""
        from packages.knowledge_store.parsers import parse_pdf

        with pytest.raises(Exception):
            parse_pdf(Path("/nonexistent/test.pdf"))


class TestParseDocx:
    """DOCX 解析器单元测试"""

    def _make_test_docx(
        self, path: Path, lines: list[str] | None = None
    ) -> None:
        """Helper: create a minimal DOCX with paragraphs."""
        from docx import Document

        doc = Document()
        if lines:
            for line in lines:
                doc.add_paragraph(line)
        doc.save(str(path))

    def test_parse_docx_returns_text(self, tmp_path: Path) -> None:
        """DOCX 解析应返回纯文本"""
        docx_path = tmp_path / "test.docx"
        self._make_test_docx(
            docx_path, ["羊肚菌产品介绍", "本品选用优质羊肚菌，口感鲜嫩"]
        )

        from packages.knowledge_store.parsers import parse_docx

        text = parse_docx(docx_path)
        assert "羊肚菌产品介绍" in text
        assert "优质羊肚菌" in text

    def test_parse_docx_empty_document(self, tmp_path: Path) -> None:
        """空 DOCX 应返回空字符串"""
        docx_path = tmp_path / "empty.docx"
        self._make_test_docx(docx_path)

        from packages.knowledge_store.parsers import parse_docx

        text = parse_docx(docx_path)
        assert text.strip() == ""

    def test_parse_docx_file_not_found(self) -> None:
        """不存在的 DOCX 应抛出异常"""
        from packages.knowledge_store.parsers import parse_docx

        with pytest.raises(Exception):
            parse_docx(Path("/nonexistent/test.docx"))


class TestParseFile:
    """parse_file 自动检测格式并路由"""

    def test_parse_file_txt(self, tmp_path: Path) -> None:
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("纯文本内容", encoding="utf-8")
        from packages.knowledge_store.parsers import parse_file

        text = parse_file(txt_path)
        assert text == "纯文本内容"

    def test_parse_file_pdf(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "test.pdf"
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_htmlbox(
            fitz.Rect(50, 50, 500, 100),
            "<p>PDF content extractable</p>",
        )
        doc.save(str(pdf_path))
        doc.close()

        from packages.knowledge_store.parsers import parse_file

        text = parse_file(pdf_path)
        assert "PDF content extractable" in text

    def test_parse_file_docx(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "test.docx"
        from docx import Document

        doc = Document()
        doc.add_paragraph("DOCX内容")
        doc.save(str(docx_path))

        from packages.knowledge_store.parsers import parse_file

        text = parse_file(docx_path)
        assert "DOCX内容" in text

    def test_parse_file_unsupported_format(self, tmp_path: Path) -> None:
        png_path = tmp_path / "test.png"
        png_path.write_text("not an image")
        from packages.knowledge_store.parsers import parse_file

        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_file(png_path)

    def test_parse_file_not_found(self) -> None:
        from packages.knowledge_store.parsers import parse_file

        with pytest.raises(Exception):
            parse_file(Path("/nonexistent/test.pdf"))


class TestParsersChineseText:
    """中文 PDF/DOCX 解析正确性（零食公司产品手册场景）"""

    def test_parse_pdf_chinese_product_brochure(self, tmp_path: Path) -> None:
        """零食公司产品手册场景：中文 PDF 解析应正确"""
        pdf_path = tmp_path / "产品手册.pdf"
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        html = (
            "<p>品牌：美味零食</p>"
            "<p>产品名称：香脆薯片</p>"
            "<p>净含量：150克</p>"
            "<p>保质期：12个月</p>"
            "<p>生产日期：见包装背面</p>"
        )
        page.insert_htmlbox(fitz.Rect(50, 50, 500, 300), html)
        doc.save(str(pdf_path))
        doc.close()

        from packages.knowledge_store.parsers import parse_pdf

        text = parse_pdf(pdf_path)
        assert "美味零食" in text
        assert "香脆薯片" in text
        assert "150克" in text
        assert "保质期" in text

    def test_parse_docx_chinese_product_brochure(self, tmp_path: Path) -> None:
        """零食公司产品手册场景：中文 DOCX 解析应正确"""
        docx_path = tmp_path / "产品手册.docx"
        from docx import Document

        doc = Document()
        doc.add_paragraph("品牌：美味零食")
        doc.add_paragraph("产品名称：香脆薯片")
        doc.add_paragraph("净含量：150克")
        doc.add_paragraph("保质期：12个月")
        doc.add_paragraph("保存方法：请置于阴凉干燥处")
        doc.save(str(docx_path))

        from packages.knowledge_store.parsers import parse_docx

        text = parse_docx(docx_path)
        assert "香脆薯片" in text
        assert "阴凉干燥处" in text
        assert "150克" in text
